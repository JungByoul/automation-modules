#!/usr/bin/env python3
"""
txt -> docx 자동 변환 폴더 감시 스크립트
인풋 폴더: Desktop/doc_input
저장 경로: 실제 WBS 폴더 구조를 실시간 스캔해 Claude가 직접 판단 및 생성
"""

import os
import sys
import json
import time
import shutil
import logging
import threading
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from anthropic import Anthropic

load_dotenv(Path(__file__).resolve().parent / ".env")
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 중복 처리 방지용 전역 상태 ──────────────────────────────────────────────
_processing_lock = threading.Lock()
_processing_files: set[str] = set()

# ─── 설정 ────────────────────────────────────────────────────────────────────
OUTPUT_BASE      = os.path.expanduser("~")
INPUT_FOLDER     = os.path.join(OUTPUT_BASE, "Desktop", "doc_input")
PROCESSED_DIR    = os.path.join(INPUT_FOLDER, "1_processed")
LOG_DIR          = os.path.join(INPUT_FOLDER, "2_log")
FINISHED_DIR     = os.path.join(INPUT_FOLDER, "3_finished")
ARCHIVE_DIR      = os.path.join(INPUT_FOLDER, "4_archive")
LOG_FILE         = str(Path(__file__).resolve().parent / "doc_watcher.log")
WBS_CONFIG_FILE  = Path(__file__).resolve().parent / "wbs_routing.json"

# ─── 로깅 ────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(LOG_FILE), exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger(__name__)

# ─── WBS 설정 로드 ───────────────────────────────────────────────────────────
_wbs_config: dict = {}

def load_wbs_config() -> dict:
    with open(WBS_CONFIG_FILE, encoding="utf-8") as f:
        return json.load(f)

# ─── 실시간 폴더 트리 스캔 ───────────────────────────────────────────────────

def _scan_dir(path: Path, lines: list, depth: int, max_depth: int):
    if depth > max_depth:
        return
    try:
        entries = sorted([e for e in path.iterdir() if e.is_dir()], key=lambda x: x.name)
        files = sorted([e for e in path.iterdir() if e.is_file() and e.suffix.lower() == ".docx"], key=lambda x: x.name)
    except PermissionError:
        return
    for entry in entries:
        lines.append("  " * depth + entry.name + "\\")
        _scan_dir(entry, lines, depth + 1, max_depth)
    for f in files:
        lines.append("  " * depth + f.name)

def scan_wbs_tree() -> str:
    """현재 실제 WBS 폴더 구조를 텍스트 트리로 반환."""
    roots = _wbs_config.get("scan_roots", ["1_Work", "2_Personal", "3_Dev", "4_Resources"])
    max_depth = _wbs_config.get("scan_depth", 3)
    lines = []
    for root_name in roots:
        root_path = Path(OUTPUT_BASE) / root_name
        if not root_path.exists():
            continue
        lines.append(root_name + "\\")
        _scan_dir(root_path, lines, depth=1, max_depth=max_depth)
    return "\n".join(lines)

# ─── Claude 시스템 프롬프트 (정적) ───────────────────────────────────────────

_SYSTEM_PROMPT = """\
당신은 문서 정리 에이전트입니다. 자유형 속기록 텍스트를 받아 구조화된 JSON으로 변환합니다.
반드시 JSON만 출력하고, 다른 텍스트는 절대 포함하지 마세요. 코드블록(```)도 쓰지 마세요.

## 메타데이터 추출
- date: YYYYMMDD 형식. 텍스트에서 찾아 변환. 못 찾으면 "미상"
- participants: 이름·직함 목록 배열. 못 찾으면 ["미상"]
- title: 핵심 주제 한 줄 요약 (20자 이내)
  - 표지에 그대로 쓰일 제목이므로 작성일자(YYMMDD/YYYYMMDD)를 title에 넣지 말 것. 날짜는 metadata.date로 별도 표기됨.
  - 주간보고류 문서는 작성일자가 아닌 보고 대상 기간 기준 "N월M주차" 형식 사용.
    좋은 예: "8월 1주차 주간업무보고" / 나쁜 예: "260807 주간업무보고", "260807 8월 1주차 주간업무보고"
- doc_type: "교육온보딩" 또는 "회의록"
  - 교육온보딩: 개념 설명, 시스템 안내, 용어 정리, 프로세스 설명 위주
  - 회의록: 논의·결정·후속조치 위주, 복수 참여자 의견 교환

## 저장 경로 결정 (save_path)
사용자 메시지 상단에 현재 실제 WBS 폴더 구조 및 각 폴더 내 기존 .docx 파일 목록이 제공됩니다.

판단 순서:
1. 내용에 가장 적합한 기존 폴더를 save_path로 지정
2. 적합한 기존 폴더가 없으면 WBS 번호 규칙에 맞게 새 폴더 경로 직접 생성:
   - 부모 폴더 번호를 그대로 이어받아 점(.) + 다음 순번
   - 같은 부모 내 마지막 번호 + 1이 새 순번
   - 폴더명은 내용 기반으로 한글/영문 혼용 가능, 언더스코어(_) 구분, 공백 금지
   - 예) 1.10 하위에 1.10.0~1.10.4가 있으면 → 1.10.5_신규폴더명
   - 예) 1_Work 하위 그룹이 1.10, 1.20까지 있으면 → 신규 그룹은 1.30_그룹명
3. 새 폴더를 만들 경우 전체 경로를 save_path에 포함 — 스크립트가 자동 생성함

save_path 형식: OUTPUT_BASE 기준 상대경로, 역슬래시(\\) 사용

## 파일명 규칙 (filename)

### ★ 기존 파일 확인 후 결정 (아래 기본 규칙보다 우선)
save_path로 결정한 폴더 안에 기존 .docx 파일이 있으면 반드시 그 패턴을 따를 것.
사람이 폴더를 오름차순 정렬로 시간순 조회한다는 전제 아래, 기존 시리즈와 나란히 섰을 때 자연스럽게 연결되는 이름 선택.

사례 1 — 순번 시리즈: 구조·순서 그대로 유지
  기존: S-OIL_mMDM_3차인수인계교육_설비등록_v0.1_20260527.docx
  잘못된 예: S-OIL_mMDM_설비MOC_인수인계교육5차_v0.1_20260609.docx  <- 순번 위치·구조 변경
  올바른 예: S-OIL_mMDM_5차인수인계교육_v0.1_20260609.docx

사례 2 — 시리즈 구조 유지: 내용 묘사로 기존 패턴을 바꾸지 말 것
  기존: S-OIL_주간보고_1회차_v0.1_20260520.docx
        S-OIL_주간보고_2회차_v0.1_20260527.docx
  잘못된 예: S-OIL_주간업무현황보고_3차_v0.1_20260603.docx  <- 구조 자체를 바꿔버림
  올바른 예: S-OIL_주간보고_3회차_v0.1_20260603.docx

사례 3 — 기존 파일 없을 때: 세부 항목 3개 이상이면 공통 주제로 압축
  잘못된 예: S-OIL_mMDM_온보딩_시스템접속_메뉴구조_기능위치등록_권한신청_v0.1_20260601.docx  <- 세부 항목 나열
  올바른 예: S-OIL_mMDM_온보딩_시스템기초_기능위치등록_v0.1_20260601.docx  <- 핵심 주제 1~2개로 압축

### 기본 규칙 (기존 파일 없을 때)
형식: [큰개념]_[중간개념]_[내용식별자]_v0.1_[날짜].docx
- 공백 금지, 구분자 _

## 출처(파일/시트) 추적 ★중요
속기록에 파일명·시트명이 언급된 경우, 해당 내용이 어느 파일·시트에서 나왔는지 반드시 섹션 content에 표시.

목차 방식 판단:
- 방식 A (파일 중심): 속기록이 파일 단위로 뚜렷이 구분되어 진행된 경우
  → 섹션 제목에 파일명 포함. 예) title="mMDM_용어집_v0.3.xlsx — 공통_용어 시트"
- 방식 B (주제 중심): 내용이 여러 파일·시트를 오가며 주제 중심으로 흐른 경우
  → 각 항목 끝에 출처 태그 명기. 예) "채번 규칙 [mMDM_용어집.xlsx > 공통_용어]"

공통 원칙:
- 출처 불명확한 내용은 [출처 미상] 표기 (임의 추정 금지)
- 한 항목이 여러 파일에 걸친 경우 모두 병기: [A.xlsx > 시트1; B.xlsx > 시트2]

## 섹션 구성 (sections 배열)

교육온보딩형:
  0번: title="후속조치_할일", content=후속 작업·질문·확인 필요 사항
  1~N번: 내용 섹션 (제목 자유롭게, 내용 기반 판단)
  마지막: title="추가메모_참고사항", content=기타 메모

회의록형:
  0번: title="회의개요", content=일시·장소·참여자·목적
  1번: title="주요논의", content=논의 내용
  2번: title="결정사항", content=결정된 사항
  3번: title="ActionItems", content=담당자·기한 포함

## 텍스트 처리 규칙
- 원문 최대한 보존
- 명백한 오타(자모 분리, 띄어쓰기) 교정. 수정 시: 수정어(원: 원문)
- 한글/영문 혼용 용어 그대로 유지 (번역 X)
- 흐름 화살표(→) 유지
- 파일명·시트명은 원문 그대로 유지 (절대 자르거나 수정·정리 금지). ~로 줄여진 부분도 그대로 옮길 것
- 문체: ~~임. ~~함. ~~있음. 명사형 단문 종결, 경어 없이

## 출력 JSON 스키마
{
  "metadata": {
    "date": "YYYYMMDD",
    "participants": ["이름1"],
    "title": "제목",
    "doc_type": "교육온보딩"
  },
  "filename": "파일명.docx",
  "save_path": "상대\\\\경로",
  "sections": [
    {"title": "섹션제목", "content": "내용 (줄바꿈은 \\n)"}
  ]
}
"""

# ─── 힌트 파싱 ───────────────────────────────────────────────────────────────

def parse_hints(content: str) -> tuple:
    hints = {}
    lines = content.split("\n")
    body_start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("#") and ":" in stripped:
            key, _, value = stripped[1:].partition(":")
            hints[key.strip()] = value.strip()
            body_start = i + 1
        else:
            break
    body = "\n".join(lines[body_start:]).strip()
    return hints, body


def get_retry_count(base_name: str) -> int:
    if not os.path.exists(LOG_DIR):
        return 0
    return len(list(Path(LOG_DIR).glob(f"{base_name}_*.txt")))


# ─── docx 생성 헬퍼 ──────────────────────────────────────────────────────────

def set_font(run, size_pt, bold=False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def set_line_spacing(para, factor=1.5):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), str(int(factor * 240)))
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def add_paragraph(doc, text="", size=11, bold=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_font(run, size, bold)
    set_line_spacing(p)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_line_spacing(p)
    return p


def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        r1 = para.add_run()
        fc1 = OxmlElement("w:fldChar")
        fc1.set(qn("w:fldCharType"), "begin")
        r1._r.append(fc1)
        r2 = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        r2._r.append(instr)
        r3 = para.add_run()
        fc2 = OxmlElement("w:fldChar")
        fc2.set(qn("w:fldCharType"), "end")
        r3._r.append(fc2)


# ─── docx 빌드 ───────────────────────────────────────────────────────────────

def build_docx(data: dict, output_path: str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    meta = data["metadata"]

    for _ in range(4):
        add_paragraph(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(meta["title"])
    set_font(title_run, 20, bold=True)
    set_line_spacing(title_p)

    for _ in range(5):
        add_paragraph(doc)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    participants = ", ".join(meta["participants"])
    info_run = info_p.add_run(f"날짜: {meta['date']}\n담당자: {participants}")
    set_font(info_run, 11)
    set_line_spacing(info_p)

    doc.add_page_break()

    add_paragraph(doc, "목차", size=13, bold=True)
    for i, sec in enumerate(data["sections"]):
        add_paragraph(doc, f"  {i}. {sec['title']}", size=11)
    add_paragraph(doc)
    doc.add_page_break()

    for i, sec in enumerate(data["sections"]):
        add_paragraph(doc, f"{i}. {sec['title']}", size=13, bold=True)
        for line in sec["content"].split("\n"):
            add_paragraph(doc, line, size=11)
        add_paragraph(doc)

    add_hr(doc)
    add_paragraph(doc, "이하 속기록", size=11, bold=True)
    add_hr(doc)
    for line in data.get("raw_content", "").split("\n"):
        add_paragraph(doc, line, size=11)

    add_page_number(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


# ─── Claude API 호출 ─────────────────────────────────────────────────────────

def call_claude(content: str, hints: dict, wbs_tree: str) -> dict:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다.")

    hint_instruction = ""
    if hints:
        lines = ["아래 힌트를 최우선으로 적용하세요. 힌트에 명시된 값은 절대 변경하지 마세요:"]
        mapping = {"경로": "save_path", "제목": "title", "타입": "doc_type", "날짜": "date", "담당자": "participants"}
        for k, v in hints.items():
            lines.append(f"- {mapping.get(k, k)}: {v}")
        hint_instruction = "\n".join(lines) + "\n\n"

    user_content = f"## 현재 WBS 폴더 구조\n{wbs_tree}\n\n"
    if hint_instruction:
        user_content += hint_instruction
    user_content += content

    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=[
            {
                "type": "text",
                "text": _SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": user_content}],
    )

    usage = response.usage
    cache_read = getattr(usage, "cache_read_input_tokens", 0) or 0
    cache_write = getattr(usage, "cache_creation_input_tokens", 0) or 0
    if cache_read > 0:
        log.info(f"캐시 적중: {cache_read} tokens read (uncached input: {usage.input_tokens})")
    elif cache_write > 0:
        log.info(f"캐시 생성: {cache_write} tokens written (uncached input: {usage.input_tokens})")

    raw = response.content[0].text.strip()
    if "```" in raw:
        for part in raw.split("```"):
            part = part.strip().lstrip("json").strip()
            if part.startswith("{"):
                raw = part
                break
    return json.loads(raw)


# ─── 로그 파일 생성 ──────────────────────────────────────────────────────────

def write_log(base_name: str, retry_count: int, data: dict, save_path: str,
              filename: str, hints: dict, success: bool, error_msg: str = ""):
    now = datetime.now()
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    log_filename = f"{base_name}_{timestamp}.txt"
    log_path = os.path.join(LOG_DIR, log_filename)

    label = "[첫처리]" if retry_count == 0 else f"[재처리({retry_count + 1}번째)]"

    meta = data.get("metadata", {}) if success else {}
    date_val = meta.get("date", "미상")
    participants = ", ".join(meta.get("participants", ["미상"]))
    title_val = meta.get("title", "미상")
    doc_type = meta.get("doc_type", "미상")

    missing = []
    if date_val == "미상": missing.append("날짜")
    if participants == "미상": missing.append("담당자")
    missing_str = ", ".join(missing) if missing else "없음"

    lines = []
    lines.append("=" * 52)
    lines.append(" txt → docx 자동 변환 로그")
    lines.append(f" {label}  |  {now.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 52)
    lines.append("")

    if success:
        lines.append("=== 변환 결과 ===")
        lines.append("상태: 성공")
        lines.append(f"파일명: {filename}")
        lines.append(f"저장경로: {save_path}")
        lines.append(f"문서타입: {doc_type}")
        lines.append("")
        lines.append("=== 추출된 메타데이터 ===")
        lines.append(f"날짜: {date_val}")
        lines.append(f"담당자: {participants}")
        lines.append(f"제목: {title_val}")
        lines.append(f"누락항목: {missing_str}")
        if missing_str != "없음":
            lines.append("  ※ 누락 항목은 원본 txt에 정보가 없어 '미상' 처리됨.")
            lines.append("    재처리 시 힌트 템플릿에 직접 입력 가능.")
    else:
        lines.append("=== 변환 결과 ===")
        lines.append("상태: 실패")
        lines.append("")
        lines.append("=== 실패 원인 ===")
        lines.append(error_msg)

    lines.append("")
    lines.append("=" * 52)
    lines.append(" 사용법")
    lines.append("=" * 52)
    lines.append("결과가 마음에 들면: 그대로 두세요.")
    if success:
        lines.append(f"  완성본 → 3_finished\\{filename}")
        lines.append(f"  WBS 저장 → {save_path}")
    lines.append("")
    lines.append("결과가 마음에 안 들면:")
    lines.append("  1. 1_processed\\ 에서 원본 txt 꺼내기")
    lines.append("  2. 파일 상단에 아래 힌트 줄 추가·수정")
    lines.append("  3. doc_input\\ 에 다시 드롭")
    lines.append("")
    lines.append("=" * 52)
    lines.append(" 재처리 힌트 템플릿 (복사 후 수정해서 사용)")
    lines.append("=" * 52)

    hint_path  = hints.get("경로",  save_path if success else "경로를_입력하세요")
    hint_title = hints.get("제목",  title_val if success else "제목을_입력하세요")
    hint_type  = hints.get("타입",  doc_type  if success else "교육온보딩 또는 회의록")
    hint_date  = hints.get("날짜",  date_val  if success else "YYYYMMDD")
    hint_part  = hints.get("담당자", participants if success else "이름")

    lines.append(f"#경로: {hint_path}")
    lines.append(f"#제목: {hint_title}")
    lines.append(f"#타입: {hint_type}")
    lines.append(f"#날짜: {hint_date}")
    lines.append(f"#담당자: {hint_part}")
    lines.append("")
    lines.append("(원본 txt 내용을 이 아래에 붙여넣기)")

    os.makedirs(LOG_DIR, exist_ok=True)
    with open(log_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    log.info(f"로그 생성: 2_log/{log_filename}")


# ─── 토스트 알림 ─────────────────────────────────────────────────────────────

def toast(title: str, message: str):
    try:
        from plyer import notification
        notification.notify(title=title, message=message, app_name="txt2docx", timeout=8)
    except Exception as e:
        log.warning(f"토스트 알림 실패: {e}")


# ─── 파일 안정화 대기 ────────────────────────────────────────────────────────

def wait_for_stable(filepath: str, max_wait: float = 5.0, interval: float = 0.3) -> bool:
    last_size = -1
    elapsed = 0.0
    while elapsed < max_wait:
        try:
            size = os.path.getsize(filepath)
        except OSError:
            return False
        if size == last_size and size > 0:
            return True
        last_size = size
        time.sleep(interval)
        elapsed += interval
    return last_size > 0


# ─── 파일 처리 ───────────────────────────────────────────────────────────────

def process_file(filepath: str):
    fname = os.path.basename(filepath)
    base_name = Path(fname).stem

    with _processing_lock:
        if filepath in _processing_files:
            log.debug(f"이미 처리 중, 스킵: {fname}")
            return
        if not os.path.exists(filepath):
            log.debug(f"파일 없음, 스킵: {fname}")
            return
        _processing_files.add(filepath)

    log.info(f"처리 시작: {fname}")

    try:
        if not wait_for_stable(filepath):
            log.warning(f"파일 안정화 실패 또는 사라짐: {fname}")
            return

        with open(filepath, encoding="utf-8") as f:
            content = f.read()
        if not content.strip():
            log.warning(f"빈 파일 무시: {fname}")
            return

        hints, body = parse_hints(content)
        retry_count = get_retry_count(base_name)
        label = f"[재처리({retry_count + 1}번째)]" if hints else "[첫처리]"
        log.info(f"{label} {fname}")

        wbs_tree = scan_wbs_tree()
        log.info(f"WBS 트리 스캔 완료")

        data = call_claude(body if body else content, hints, wbs_tree)
        data["raw_content"] = body if body else content

        save_dir  = os.path.join(OUTPUT_BASE, data["save_path"])
        save_path = os.path.join(save_dir, data["filename"])
        build_docx(data, save_path)
        log.info(f"저장 완료: {save_path}")

        os.makedirs(FINISHED_DIR, exist_ok=True)
        shutil.copy2(save_path, os.path.join(FINISHED_DIR, data["filename"]))

        os.makedirs(ARCHIVE_DIR, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy2(save_path, os.path.join(ARCHIVE_DIR, f"{ts}_{data['filename']}"))

        os.makedirs(PROCESSED_DIR, exist_ok=True)
        shutil.move(filepath, os.path.join(PROCESSED_DIR, fname))

        write_log(base_name, retry_count, data, data["save_path"],
                  data["filename"], hints, success=True)

        toast(f"변환 완료 {label}", f"{data['filename']}\n→ {data['save_path']}")
        log.info(f"완료: {fname}")

    except Exception as e:
        err_msg = str(e)
        log.error(f"오류 ({fname}): {err_msg}", exc_info=True)
        write_log(base_name, get_retry_count(base_name), {}, "", "", {}, success=False, error_msg=err_msg)
        toast("변환 실패", f"{fname}\n{err_msg[:80]}")
    finally:
        with _processing_lock:
            _processing_files.discard(filepath)


# ─── 감시 루프 ───────────────────────────────────────────────────────────────

class TxtHandler(FileSystemEventHandler):
    def _maybe_process(self, event):
        if event.is_directory:
            return
        if not event.src_path.lower().endswith(".txt"):
            return
        process_file(event.src_path)

    def on_created(self, event):
        self._maybe_process(event)

    def on_modified(self, event):
        self._maybe_process(event)


def main():
    global _wbs_config
    _wbs_config = load_wbs_config()
    log.info(f"WBS 스캔 루트: {_wbs_config.get('scan_roots')}")

    for d in [INPUT_FOLDER, PROCESSED_DIR, LOG_DIR, FINISHED_DIR, ARCHIVE_DIR]:
        os.makedirs(d, exist_ok=True)

    log.info(f"감시 시작: {INPUT_FOLDER}")
    log.info("txt 파일을 인풋 폴더에 넣으면 자동 변환됩니다.")

    for f in Path(INPUT_FOLDER).glob("*.txt"):
        process_file(str(f))

    observer = Observer()
    observer.schedule(TxtHandler(), INPUT_FOLDER, recursive=False)
    observer.start()
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        log.info("감시 종료.")
        observer.stop()
    observer.join()


if __name__ == "__main__":
    main()
